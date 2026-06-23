from google import genai
from google.genai import types
from django.conf import settings
import PyPDF2
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import tempfile
import os
import re

WEB_SEARCH_AVAILABLE = False

try:

    pytesseract.pytesseract.pytesseract_cmd = 'tesseract'
except:

    try:
        pytesseract.pytesseract.pytesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    except:

        print("Note: Tesseract not found. Gemini Vision API will be used for image/image-PDF processing.")

google_api_key = getattr(settings, 'GOOGLE_API_KEY', None) or os.getenv('GOOGLE_API_KEY')
if not google_api_key:
    print("WARNING: GOOGLE_API_KEY is not set. AI features will not work until it is configured.")
    google_api_key = 'not-configured'

client = genai.Client(api_key=google_api_key)

# Monkeypatch generate_content to handle rate limits, transient 503 errors,
# and fall back to stable models under high demand.
_original_generate_content = client.models.generate_content

def generate_content_with_retry(*args, **kwargs):
    import time
    max_retries = 3
    delay = 1
    last_error = None
    
    # Extract and track the model being used
    model = kwargs.get('model', 'gemini-2.5-flash')
    current_model = model
    
    for attempt in range(max_retries):
        try:
            kwargs['model'] = current_model
            return _original_generate_content(*args, **kwargs)
        except Exception as e:
            last_error = e
            err_str = str(e).upper()
            if attempt < max_retries - 1 and any(token in err_str for token in ("503", "UNAVAILABLE", "429", "RESOURCE_EXHAUSTED", "LIMIT")):
                print(f"Gemini API returned error ({e}) for model {current_model}. Retrying in {delay}s...")
                time.sleep(delay)
                delay *= 2
                # Fallback to stable models if the default is overloaded
                if attempt == max_retries - 2:
                    if current_model == 'gemini-2.5-flash':
                        current_model = 'gemini-2.0-flash'
                    elif current_model == 'gemini-2.0-flash':
                        current_model = 'gemini-1.5-flash'
            else:
                raise e
    if last_error:
        raise last_error

client.models.generate_content = generate_content_with_retry

try:
    import sympy as sp
    SYMPY_AVAILABLE = True
except ImportError:
    SYMPY_AVAILABLE = False
    print("Warning: SymPy not installed. pip install sympy to enable verified math solving.")

try:
    import scipy
    import scipy.optimize
    import scipy.integrate
    import scipy.linalg
    import scipy.stats
    import scipy.constants
    import numpy as np
    SCIPY_AVAILABLE = True
except ImportError:
    SCIPY_AVAILABLE = False
    print("Warning: SciPy/NumPy not installed. pip install scipy numpy to enable physics/engineering solving.")

try:
    import pint
    _ureg = pint.UnitRegistry()
    PINT_AVAILABLE = True
except ImportError:
    PINT_AVAILABLE = False
    print("Warning: Pint not installed. pip install pint to enable unit-aware calculations.")

try:
    import chempy
    CHEMPY_AVAILABLE = True
except ImportError:
    CHEMPY_AVAILABLE = False
    print("Warning: ChemPy not installed. pip install chempy to enable chemistry solving.")

_MATH_KEYWORDS = {
    'solve', 'calculate', 'compute', 'evaluate', 'simplify', 'differentiate',
    'integrate', 'derivative', 'integral', 'limit', 'expand', 'factor',
    'prove', 'determinant', 'eigenvalue', 'eigenvalues', 'matrix',
    'series', 'taylor', 'fourier', 'laplace', 'equation', 'polynomial',
    'roots', 'zeros', 'gradient', 'divergence', 'curl',
}

def _has_whole_word(text: str, keywords: set) -> bool:
    lower_text = text.lower()
    words = set(re.findall(r'\b\w+\b', lower_text))
    for kw in keywords:
        if ' ' in kw:
            if kw in lower_text:
                return True
        else:
            if kw in words:
                return True
    return False

def is_math_computation_problem(message: str) -> bool:
    if _has_whole_word(message, _MATH_KEYWORDS):
        return True
    if re.search(r'[\^=√∫∑∏]|d/dx|\bx\b.*=|\bf\(x\)', message.lower()):
        return True
    return False

_PHYSICS_ENGINEERING_KEYWORDS = {

    'velocity', 'acceleration', 'force', 'momentum', 'impulse', 'torque', 'angular',
    'friction', 'gravity', 'gravitational', 'kinematics', 'dynamics', 'statics',
    'projectile', 'trajectory', 'centripetal', 'centrifugal',

    'kinetic energy', 'potential energy', 'joule', 'watt',

    'temperature', 'heat', 'entropy', 'pressure', 'ideal gas',
    'thermodynamics', 'carnot', 'kelvin', 'celsius', 'boyle', 'charles',

    'voltage', 'current', 'resistance', 'capacitance', 'inductance', 'impedance',
    'circuit', 'ohm', 'ampere', 'farad', 'henry', 'coulomb',
    'electric field', 'magnetic field', 'flux', 'electromagnetic',

    'wavelength', 'frequency', 'amplitude', 'oscillation', 'resonance',
    'refraction', 'reflection', 'diffraction', 'interference', 'optics',

    'photon', 'quantum', 'relativity', 'nuclear',

    'stress', 'strain', 'modulus', 'elasticity', 'shear', 'bending',
    'resistor', 'capacitor', 'inductor', 'transistor', 'diode',

    'newton', 'pascal', 'hertz', 'tesla', 'siemens',
}

def is_physics_engineering_problem(message: str) -> bool:
    if _has_whole_word(message, _PHYSICS_ENGINEERING_KEYWORDS):
        return True
    if re.search(r'\d+\s*(m/s|km/h|\bkg\b|m/s²|\bN\b|\bPa\b|\bJ\b|\bW\b|\bV\b|\bA\b|\bHz\b|rad/s)', message):
        return True
    return False

_CHEMISTRY_KEYWORDS = {
    'mole', 'molar', 'molarity', 'molality', 'concentration',
    'stoichiometry', 'stoichiometric',
    'acid', 'base', 'ph', 'buffer', 'titration', 'neutralization',
    'oxidation', 'reduction', 'redox',
    'gibbs', 'thermochemistry', 'hess',
    'equilibrium', 'rate constant', 'reaction rate',
    'periodic table', 'atomic number', 'atomic mass', 'isotope', 'atomic weight',
    'organic chemistry', 'functional group', 'hydrocarbon', 'polymer',
    'reactant', 'product', 'yield', 'catalyst',
    'electron configuration', 'valence', 'orbital', 'hybridization',
    'electrochemistry', 'cell potential', 'faraday',
    'avogadro', 'dalton', 'ideal gas law',
    'solubility', 'precipitate', 'dissolution',
    'chemical equation', 'balance the equation', 'balanced equation',
}

def is_chemistry_problem(message: str) -> bool:
    if _has_whole_word(message, _CHEMISTRY_KEYWORDS):
        return True
    formula_match = re.search(r'\b(?:[A-Z][a-z]?\d*){2,}\b', message)
    if formula_match:
        token = formula_match.group(0)
        has_digit = any(ch.isdigit() for ch in token)
        has_lower = any(ch.islower() for ch in token)
        if has_digit or has_lower:
            return True
    return False

def solve_with_sympy(problem_description: str):
    """
    Ask Gemini to generate SymPy code for a math problem, execute it safely,
    and return (plain_result, latex_result).  Returns (None, None) on failure.
    """
    if not SYMPY_AVAILABLE:
        return None, None
    try:
        code_prompt = (
            "Write self-contained Python code using SymPy to solve the following math problem.\n"
            "Rules:\n"
            "- Import sympy at the top.\n"
            "- Store the final answer in a variable called `result`.\n"
            "- Store the LaTeX string of the answer in `result_latex` using sp.latex().\n"
            "- Output ONLY the raw Python code — no markdown fences, no prose.\n\n"
            f"Problem: {problem_description}"
        )
        code_response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=code_prompt
        )
        code = code_response.text.strip()

        code = re.sub(r'^```(?:python)?\s*', '', code, flags=re.MULTILINE)
        code = re.sub(r'```\s*$', '', code, flags=re.MULTILINE)
        code = code.strip()

        namespace = {'__builtins__': {
            'print': print, 'range': range, 'len': len, 'list': list,
            'dict': dict, 'set': set, 'tuple': tuple, 'int': int,
            'float': float, 'str': str, 'bool': bool, 'abs': abs,
            'round': round, 'enumerate': enumerate, 'zip': zip,
            '__import__': __import__,
        }}
        exec(code, namespace)

        result = namespace.get('result')
        result_latex = namespace.get('result_latex', '')
        if result is None:
            return None, None
        if not result_latex:
            result_latex = sp.latex(result)
        return str(result), str(result_latex)
    except Exception as exc:
        print(f"SymPy solver error (non-blocking): {exc}")
        return None, None

def solve_with_scipy_pint(problem_description: str):
    """
    Ask Gemini to generate SciPy/NumPy/Pint code for a physics or engineering
    problem, execute it safely, and return a plain-text result string.
    Returns None on failure.
    """
    if not SCIPY_AVAILABLE:
        return None
    try:
        code_prompt = (
            "Write self-contained Python code using SciPy, NumPy, and optionally Pint "
            "to solve the following physics or engineering problem.\n"
            "Rules:\n"
            "- Import scipy, numpy as np, and pint at the top as needed.\n"
            "- Store a human-readable string of the final answer in a variable called `result`.\n"
            "- Include units in `result` where relevant (e.g. '9.81 m/s²', '220 V').\n"
            "- Output ONLY the raw Python code — no markdown fences, no prose.\n\n"
            f"Problem: {problem_description}"
        )
        code_response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=code_prompt
        )
        code = code_response.text.strip()
        code = re.sub(r'^```(?:python)?\s*', '', code, flags=re.MULTILINE)
        code = re.sub(r'```\s*$', '', code, flags=re.MULTILINE)
        code = code.strip()

        namespace = {
            '__builtins__': {
                'print': print, 'range': range, 'len': len, 'list': list,
                'dict': dict, 'set': set, 'tuple': tuple, 'int': int,
                'float': float, 'str': str, 'bool': bool, 'abs': abs,
                'round': round, 'enumerate': enumerate, 'zip': zip,
                'map': map, 'min': min, 'max': max, 'sum': sum,
                '__import__': __import__,
            },
            'scipy': scipy,
            'np': np,
            'numpy': np,
        }
        if PINT_AVAILABLE:
            namespace['pint'] = pint
            namespace['ureg'] = pint.UnitRegistry()

        exec(code, namespace)
        result = namespace.get('result')
        return str(result) if result is not None else None
    except Exception as exc:
        print(f"SciPy/Pint solver error (non-blocking): {exc}")
        return None

def solve_chemistry(problem_description: str):
    """
    Ask Gemini to generate ChemPy/SciPy/NumPy code for a chemistry problem,
    execute it safely, and return a plain-text result string.
    Returns None on failure.
    """
    if not CHEMPY_AVAILABLE and not SCIPY_AVAILABLE:
        return None
    try:
        code_prompt = (
            "Write self-contained Python code using ChemPy and/or SciPy/NumPy "
            "to solve the following chemistry problem.\n"
            "Rules:\n"
            "- Import chempy, scipy, numpy as np at the top as needed.\n"
            "- Store a human-readable string summary of the final answer in a variable called `result`.\n"
            "- For stoichiometry: include molar masses, moles, and amounts with units.\n"
            "- For equilibrium: include equilibrium concentrations and Kc/Kp.\n"
            "- For thermochemistry: include ΔH, ΔS, ΔG with units.\n"
            "- For pH: include the pH value and relevant ion concentrations.\n"
            "- Output ONLY the raw Python code — no markdown fences, no prose.\n\n"
            f"Problem: {problem_description}"
        )
        code_response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=code_prompt
        )
        code = code_response.text.strip()
        code = re.sub(r'^```(?:python)?\s*', '', code, flags=re.MULTILINE)
        code = re.sub(r'```\s*$', '', code, flags=re.MULTILINE)
        code = code.strip()

        namespace = {
            '__builtins__': {
                'print': print, 'range': range, 'len': len, 'list': list,
                'dict': dict, 'set': set, 'tuple': tuple, 'int': int,
                'float': float, 'str': str, 'bool': bool, 'abs': abs,
                'round': round, 'enumerate': enumerate, 'zip': zip,
                'map': map, 'min': min, 'max': max, 'sum': sum, 'sorted': sorted,
                '__import__': __import__,
            },
        }
        if SCIPY_AVAILABLE:
            namespace['scipy'] = scipy
            namespace['np'] = np
            namespace['numpy'] = np
        if CHEMPY_AVAILABLE:
            namespace['chempy'] = chempy

        exec(code, namespace)
        result = namespace.get('result')
        return str(result) if result is not None else None
    except Exception as exc:
        print(f"ChemPy solver error (non-blocking): {exc}")
        return None

_SCANNER_WATERMARKS = [
    'camscanner', 'adobe scan', 'microsoft lens', 'genius scan', 'anyscanner',
    'tiny scanner', 'turbo scan', 'scanbot', 'docscanner',
    'scan with', 'scanned by', 'scanned with',
]

def _strip_scanner_watermark_noise(text: str) -> str:
    """Remove common scanner watermark lines from OCR output."""
    if not text:
        return ""

    cleaned_lines = []
    for line in text.splitlines():
        compact = re.sub(r'\s+', ' ', line).strip()
        if not compact:
            cleaned_lines.append(line)
            continue

        lowered = compact.lower()
        lowered_no_space = re.sub(r'\s+', '', lowered)

        if any(wm.replace(' ', '') in lowered_no_space for wm in _SCANNER_WATERMARKS):
            continue
        if 'scanner' in lowered and ('http' in lowered or '.com' in lowered or 'www.' in lowered):
            continue

        cleaned_lines.append(line)

    cleaned_text = "\n".join(cleaned_lines)

    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text).strip()
    return cleaned_text

def _is_meaningful_text(text: str, min_chars: int = 120) -> bool:
    """
    Return True only if the extracted text contains substantial real content
    after removing known scanner watermark noise.
    """
    cleaned = _strip_scanner_watermark_noise(text).lower()
    for watermark in _SCANNER_WATERMARKS:
        cleaned = cleaned.replace(watermark, '')

    import re
    cleaned = re.sub(r'[\s\W\d]+', ' ', cleaned).strip()
    return len(cleaned) >= min_chars

def extract_text_from_pdf(pdf_path):
    """Extract text content from PDF file with fallback to Gemini vision for image-based PDFs"""
    text = ""
    try:

        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"

        if text.strip() and _is_meaningful_text(text):
            return text

        print("Text-based extraction yielded no meaningful content, falling back to Gemini Vision...")
        text = extract_text_from_pdf_with_gemini_vision(pdf_path)

    except Exception as e:

        print(f"PyPDF2 extraction error, falling back to Gemini vision: {str(e)}")
        try:
            text = extract_text_from_pdf_with_gemini_vision(pdf_path)
        except Exception as e2:
            raise Exception(f"Failed to extract PDF text: {str(e2)}")

    return text if text.strip() else "Unable to extract text from this PDF."

def extract_text_from_pdf_with_gemini_vision(pdf_path):
    """
    Extract text from image-based PDFs (including scanned / handwritten pages)
    using Gemini's vision API.
    Speed/quality balance:
    - 260 DPI default for better small-text recovery on blurry scans
    - Images resized to max 1 600 px wide before encoding
    - JPEG encoding (5-10× smaller than PNG)
    - All pages processed IN PARALLEL via ThreadPoolExecutor
    Quality:
    - Adaptive preprocessing with conditional retry variants for dark/low-contrast pages
    - Gemini prompt explicitly ignores scanner watermarks (CamScanner, etc.)
    """
    import base64
    from PIL import ImageEnhance, ImageOps, ImageFilter
    from concurrent.futures import ThreadPoolExecutor, as_completed

    poppler_path = r'C:\Users\HomePC\Downloads\poppler\poppler-25.12.0\Library\bin'

    try:
        try:
            if os.path.exists(poppler_path):
                images = convert_from_path(
                    pdf_path, first_page=1, last_page=20,
                    dpi=260, poppler_path=poppler_path
                )
            else:
                images = convert_from_path(pdf_path, first_page=1, last_page=20, dpi=260)
        except Exception as e:
            print(f"Poppler path failed, trying system poppler: {e}")
            images = convert_from_path(pdf_path, first_page=1, last_page=20, dpi=260)

    except Exception as e:
        raise Exception(f"Failed to convert PDF pages to images: {str(e)}")

    PAGE_PROMPT = (
        "You are reading a scanned document page.\n"
        "Your task: extract ONLY the actual document content — "
        "handwritten notes, printed text, diagrams, tables, equations, "
        "and any legible writing made by the document author.\n\n"
        "IMPORTANT RULES:\n"
        "1. IGNORE all scanner / app watermarks, logos, and branding. "
        "This includes 'CamScanner', 'Adobe Scan', 'Microsoft Lens', "
        "'Genius Scan', 'AnyScanner', any app name, website URL, or promotional text "
        "added by a scanning app — do NOT transcribe these.\n"
        "2. If the page contains handwriting, transcribe it faithfully, "
        "preserving line breaks, numbering, and structure.\n"
        "3. If text is partially illegible, give your best reading and "
        "mark uncertain words with [?].\n"
        "4. Preserve the original layout: headings, bullet points, "
        "numbered lists, tables, and paragraph breaks.\n"
        "5. Return ONLY the transcribed text — no commentary or explanations."
    )

    def _text_quality_score(text: str) -> int:
        """Heuristic score to decide whether OCR output is usable."""
        if not text:
            return 0

        cleaned = text.strip()
        if not cleaned:
            return 0

        alnum_chars = sum(1 for c in cleaned if c.isalnum())
        total_chars = len(cleaned)
        words = [w for w in re.split(r'\s+', cleaned) if w]
        long_words = sum(1 for w in words if len(w) >= 3)
        line_count = len([line for line in cleaned.splitlines() if line.strip()])

        score = 0
        score += min(total_chars, 1200) // 8
        score += min(alnum_chars, 1000) // 10
        score += min(long_words, 200)
        score += min(line_count, 60) * 2
        return score

    def _build_preprocessed_variants(page_img):
        """Generate ordered image variants for robust OCR on difficult scans."""
        variants = []

        v1 = page_img.convert('RGB')
        v1 = ImageEnhance.Contrast(v1).enhance(1.9)
        v1 = ImageEnhance.Sharpness(v1).enhance(2.2)
        v1 = ImageEnhance.Brightness(v1).enhance(1.08)
        variants.append(v1)

        v2 = page_img.convert('L')
        v2 = ImageOps.autocontrast(v2, cutoff=1)
        v2 = ImageEnhance.Brightness(v2).enhance(1.35)
        v2 = ImageEnhance.Contrast(v2).enhance(2.0)
        v2 = v2.filter(ImageFilter.UnsharpMask(radius=1.8, percent=170, threshold=2))
        variants.append(v2.convert('RGB'))

        v3 = page_img.convert('L').filter(ImageFilter.MedianFilter(size=3))
        v3 = ImageOps.autocontrast(v3, cutoff=2)
        v3 = v3.point(lambda p: 255 if p > 150 else 0)
        variants.append(v3.convert('RGB'))

        return variants

    def _extract_with_gemini(prepared_img):
        """Encode one preprocessed image and ask Gemini for transcription."""
        tmp_path = None
        try:
            max_width = 1600
            if prepared_img.width > max_width:
                ratio = max_width / prepared_img.width
                prepared_img = prepared_img.resize(
                    (max_width, int(prepared_img.height * ratio)),
                    resample=Image.LANCZOS
                )

            with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp:
                prepared_img.save(tmp, format='JPEG', quality=92, optimize=True)
                tmp_path = tmp.name

            with open(tmp_path, 'rb') as f:
                img_bytes = f.read()

            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[
                    PAGE_PROMPT,
                    types.Part.from_bytes(
                        data=img_bytes,
                        mime_type="image/jpeg",
                    )
                ]
            )
            return response.text.strip() if response and response.text else ""
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def process_page(args):
        """Preprocess one page image and call Gemini. Returns (idx, text)."""
        idx, img = args
        try:
            variants = _build_preprocessed_variants(img)

            best_text = _extract_with_gemini(variants[0])
            best_score = _text_quality_score(best_text)

            if best_score < 65:
                for variant in variants[1:]:
                    candidate_text = _extract_with_gemini(variant)
                    candidate_score = _text_quality_score(candidate_text)
                    if candidate_score > best_score:
                        best_text = candidate_text
                        best_score = candidate_score

                    if best_score >= 120:
                        break

            best_text = _strip_scanner_watermark_noise(best_text)
            return (idx, best_text)

        except Exception as e:
            print(f"Error processing page {idx + 1} with Gemini: {e}")
            return (idx, "")

    results = {}
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = {executor.submit(process_page, (idx, img)): idx
                   for idx, img in enumerate(images)}
        for future in as_completed(futures):
            idx, page_text = future.result()
            if page_text:
                results[idx] = page_text

    full_text = "\n\n".join(results[i] for i in sorted(results))
    full_text = _strip_scanner_watermark_noise(full_text)
    return full_text if full_text.strip() else "No readable text found in this PDF."

def is_tesseract_available():
    """Check if tesseract is installed and available"""
    try:
        pytesseract.get_tesseract_version()
        return True
    except:
        return False

def extract_text_from_pdf_with_ocr(pdf_path):
    """DEPRECATED: Use extract_text_from_pdf_with_gemini_vision instead"""
    return extract_text_from_pdf_with_gemini_vision(pdf_path)

def extract_text_from_image(image_path):
    """Extract text from image using Gemini's vision API (primary) or OCR fallback"""
    try:

        with open(image_path, 'rb') as f:
            img_bytes = f.read()
        
        image_type = "image/jpeg"
        if image_path.lower().endswith('.png'):
            image_type = "image/png"
        elif image_path.lower().endswith('.gif'):
            image_type = "image/gif"
        elif image_path.lower().endswith('.webp'):
            image_type = "image/webp"
        
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                (
                    "You are reading a scanned or photographed document/image.\n"
                    "Your task: extract ONLY the actual content created by the document author — "
                    "handwritten text, printed text, diagrams, tables, equations, labels, and captions.\n\n"
                    "IMPORTANT RULES:\n"
                    "1. IGNORE all scanner / app watermarks, logos, and branding. "
                    "This includes 'CamScanner', 'Adobe Scan', 'Microsoft Lens', "
                    "'AnyScanner', "
                    "any app name, website URL, or promotional overlay added by a scanning app.\n"
                    "2. Transcribe handwriting faithfully, preserving the original line breaks and structure.\n"
                    "3. If text is partially illegible, give your best reading and mark uncertain words with [?].\n"
                    "4. If it contains a diagram or chart, describe its structure and all labelled values.\n"
                    "5. Return only the transcribed content — no commentary or explanations."
                ),
                types.Part.from_bytes(
                    data=img_bytes,
                    mime_type=image_type,
                )
            ]
        )
        
        if response.text.strip():
            return _strip_scanner_watermark_noise(response.text)
        else:
            return "No readable text found in this image."
            
    except Exception as e:
        print(f"Gemini vision failed: {e}")

        try:
            if is_tesseract_available():
                image = Image.open(image_path)
                text = pytesseract.image_to_string(image)
                if text.strip():
                    return text
        except Exception as ocr_error:
            print(f"OCR fallback also failed: {ocr_error}")
        
        return "Unable to extract text from this image. Try asking questions about it and I'll help analyze it!"

def extract_text_from_word(doc_path):
    """Extract text from Word document (.docx or .doc)"""
    try:
        doc = Document(doc_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text += " | ".join(row_text) + "\n"
        
        if not text.strip():
            return "No readable text found in this Word document."
        
        return text
    except Exception as e:
        return f"Failed to extract text from Word document: {str(e)}"

def extract_text_from_plain_text(text_path):
    """Extract text from a plain-text file."""
    try:
        with open(text_path, 'r', encoding='utf-8') as text_file:
            content = text_file.read()
        return content if content.strip() else "No readable text found in this text document."
    except UnicodeDecodeError:
        with open(text_path, 'r', encoding='latin-1') as text_file:
            content = text_file.read()
        return content if content.strip() else "No readable text found in this text document."
    except Exception as e:
        return f"Failed to extract text from text document: {str(e)}"

def summarize_pdf(pdf_path, user_instruction=None):
    """
    Summarize PDF content with structured formatting using Google Gemini
    """
    try:
        pdf_text = extract_text_from_pdf(pdf_path)
        
        if not pdf_text.strip():
            return "Unable to extract text from this PDF. The document may be image-based or encrypted."
        
        pdf_text = pdf_text[:15000]
        
        prompt = f"""You are LearnBuddy. Analyze this material and provide a STYLED summary.

### FORMATTING RULES:
1. Use ## for Section Headers.
2. Use * for Bullet Points.
3. Use --- to separate sections.
4. Always put a double line break between paragraphs.

Document Content:
{pdf_text}

Please provide:
## Overview
(2-3 sentences about the main topic)

---
## Key Concepts
* (Concept 1)
* (Concept 2)

---
## Learning Objectives
* (Goal 1)

---
## Notable Facts
* (Fact 1)

If this contains religious content, highlight it warmly. Format in a friendly, helpful tone.
In the course of summarizing documents, do not give the same response as the general response. Give a more clear, precise and concise explanation with more detailed explanation about the document's content."""

        if user_instruction:
            prompt += f"\n\n**User's specific request:** {user_instruction}\nMake sure to address this specific request directly in your response."

        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        result = response.text
        print(f"Successfully summarized PDF using Google Gemini 2.5 Flash")
        return result
        
    except Exception as e:
        return f"I processed the PDF, but encountered an issue generating a detailed summary. Error: {str(e)}"

def summarize_image(image_path, user_instruction=None):
    """
    Extract and analyze text from image with structured formatting using Google Gemini
    """
    try:
        image_text = extract_text_from_image(image_path)
        
        if not image_text.strip():
            return "Unable to extract text from this image. The image may be too blurry or contain no readable text."
        
        image_text = image_text[:8000]
        
        prompt = f"""You are LearnBuddy. Analyze this text extracted from an image and provide a STYLED summary.

### FORMATTING RULES:
1. Use ## for Section Headers.
2. Use * for Bullet Points.
3. Use --- to separate sections.
4. Always put a double line break between paragraphs.

Extracted Text from Image:
{image_text}

Please provide:
## Overview
(2-3 sentences about the main content)

---
## Key Points
* (Point 1)
* (Point 2)

---
## Notable Information
* (Info 1)

Format in a friendly, helpful tone. Be clear and precise in your explanation."""

        if user_instruction:
            prompt += f"\n\n**User's specific request:** {user_instruction}\nMake sure to address this specific request directly in your response."

        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        result = response.text
        print(f"Successfully analyzed image using Google Gemini 2.5 Flash")
        return result
        
    except Exception as e:
        return f"I processed the image, but encountered an issue generating a summary. Error: {str(e)}"

def summarize_document(doc_path, user_instruction=None):
    """
    Extract and summarize text from document files with structured formatting
    """
    try:
        if doc_path.lower().endswith('.txt'):
            doc_text = extract_text_from_plain_text(doc_path)
        else:
            doc_text = extract_text_from_word(doc_path)
        
        if not doc_text.strip() or "Failed to extract" in doc_text:
            return doc_text if doc_text else "Unable to extract text from this document."
        
        doc_text = doc_text[:8000]
        
        prompt = f"""You are LearnBuddy. Analyze this text extracted from a Word document and provide a STYLED summary.

### FORMATTING RULES:
1. Use ## for Section Headers.
2. Use * for Bullet Points.
3. Use --- to separate sections.
4. Always put a double line break between paragraphs.

Document Content:
{doc_text}

Please provide:
## Overview
(2-3 sentences about the document)

---
## Key Concepts
* (Concept 1)
* (Concept 2)

---
## Important Topics
* (Topic 1)

---
## Key Takeaways
* (Takeaway 1)

Format in a friendly, helpful tone. Be clear and precise in your explanation."""

        if user_instruction:
            prompt += f"\n\n**User's specific request:** {user_instruction}\nMake sure to address this specific request directly in your response."

        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        result = response.text
        print(f"Successfully summarized document using Google Gemini 2.5 Flash")
        return result
        
    except Exception as e:
        return f"I processed the document, but encountered an issue generating a summary. Error: {str(e)}"

def generate_session_title(user_message: str, assistant_response: str) -> str:
    """
    Generate a concise 4-6 word title that captures what the conversation is about.
    Called once after the first exchange so the title reflects the actual topic.
    """
    try:
        prompt = (
            "Based on this conversation exchange, write a short chat title of 4 to 6 words max. "
            "The title should capture the main topic clearly. "
            "Do NOT use quotes, punctuation, or any prefix like 'Title:'. Just the plain words.\n\n"
            f"User: {user_message[:300]}\n"
            f"Assistant: {assistant_response[:300]}"
        )
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        title = response.text.strip().strip('"\'')

        return title[:80] if title else ''
    except Exception:
        return ''

def ask_buddy(user_message, conversation_history=None, material_context=None, 
              system_context=None, is_religion_topic=False, file=None):
    """
    Get AI response with improved layout using Google Gemini
    Includes real-time web search for current events/news questions
    """
    try:
        system_message = """You are LearnBuddy. Your personality:
1. RELIGIOUS TOPICS: Give clear and logical answers, backed up by the scripture/Bible and Quran to whatever religious question. either christian, islamic, or whatever religion even down to atheism and buddhism and the rest.
2. EDUCATIONAL CONTENT: Break down complex topics using bullet points and headers.
3. GENERAL TONE: Friendly and organized. Always use double line breaks between ideas.
4. MATHEMATICS: For ALL mathematical content you MUST:
   - Use LaTeX notation exclusively — never write math in plain words.
   - Inline expressions: wrap with $...$ (e.g., $x^2 + y^2 = z^2$).
   - Display / block equations: wrap with $$...$$ on its own line.
   - Use full LaTeX syntax: \\frac{a}{b}, \\int_{a}^{b} f(x)\\,dx, \\sum_{n=0}^{\\infty}, \\sqrt{x}, \\lim_{x \\to 0}, etc.
   - Greek letters via LaTeX: \\alpha, \\beta, \\pi, \\theta, \\Delta, \\Sigma, etc.
   - NEVER write "integral from a to b" — write $$\\int_a^b f(x)\\,dx$$ instead.
   - NEVER use Unicode superscripts (², ³) — use $x^2$, $x^3$.
   - Show step-by-step solutions with each step wrapped in LaTeX.
   - If a SYMPY / SCIPY / CHEMPY VERIFIED RESULT is provided below, use that exact answer.
5. INTERNAL CONTEXT RULES: The prompt may contain blocks labelled 'SYMPY VERIFIED RESULT', 'SCIPY/PINT NUMERICAL RESULT', 'CHEMPY VERIFIED RESULT', 'STUDY MATERIAL CONTEXT', or 'Previous conversation'. These are SILENT internal hints for you only. NEVER copy, quote, repeat, or reference these block labels or their raw content in your reply. Use the values to inform your answer, but write naturally as if you computed them yourself."""

        if system_context:
            system_message = system_context
        
        current_event_info = ""

        stem_context = ""

        if is_math_computation_problem(user_message):
            sympy_plain, sympy_latex = solve_with_sympy(user_message)
            if sympy_plain:
                stem_context += (
                    f"SYMPY VERIFIED RESULT (computed symbolically — use this exact answer):\n"
                    f"  Plain:  {sympy_plain}\n"
                    f"  LaTeX:  ${sympy_latex}$\n\n"
                )
                print(f"SymPy result for '{user_message[:60]}': {sympy_plain}")

        if is_physics_engineering_problem(user_message):
            sci_result = solve_with_scipy_pint(user_message)
            if sci_result:
                stem_context += f"SCIPY/PINT NUMERICAL RESULT (use this exact value):\n  {sci_result}\n\n"
                print(f"SciPy result for '{user_message[:60]}': {sci_result}")

        if is_chemistry_problem(user_message):
            chem_result = solve_chemistry(user_message)
            if chem_result:
                stem_context += f"CHEMPY VERIFIED RESULT (use this exact value):\n  {chem_result}\n\n"
                print(f"ChemPy result for '{user_message[:60]}': {chem_result}")
        
        conversation_text = ""
        if conversation_history:
            for msg in conversation_history[-12:]:
                role = msg.get('role', 'user')
                if 'parts' in msg:
                    content = msg['parts'][0] if msg['parts'] else ""
                elif 'text' in msg:
                    content = msg['text']
                elif 'content' in msg:
                    content = msg['content']
                else:
                    content = str(msg)
                
                if role == 'model':
                    role = 'Assistant'
                elif role == 'assistant':
                    role = 'Assistant'
                else:
                    role = 'User'
                
                if content:
                    conversation_text += f"{role}: {str(content)[:1000]}\n\n"
        
        full_prompt = system_message + "\n\n"
        
        if conversation_text:
            full_prompt += "Previous conversation:\n" + conversation_text + "\n"
        
        if material_context:
            full_prompt += f"STUDY MATERIAL CONTEXT:\n{material_context[:4000]}\n\n"
        
        if current_event_info:
            full_prompt += current_event_info + "\n"

        if stem_context:
            full_prompt += stem_context + "\n"

        if is_religion_topic:
            full_prompt += "The user is asking about religious topics. Respond with warmth, scriptural and logical references using clear headers.\n\n"
        
        full_prompt += f"User: {user_message}\nAssistant:"
        
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=full_prompt
        )
        result = response.text

        _internal_block_pattern = re.compile(
            r'(?:SYMPY VERIFIED RESULT|SCIPY/PINT NUMERICAL RESULT|CHEMPY VERIFIED RESULT|STUDY MATERIAL CONTEXT)'
            r'[^\n]*\n(?:\s+[^\n]+\n)*',
            re.IGNORECASE
        )
        result = _internal_block_pattern.sub('', result).strip()

        return result
        
    except Exception as e:
        if is_religion_topic:
            return "I'm experiencing a technical issue. Please share a specific verse you'd like to discuss, or feel free to rephrase your question."
        return f"I'm here to help, but I encountered a technical issue. (Error: {str(e)})"

from pydantic import BaseModel
import json

class FlashcardItem(BaseModel):
    front: str
    back: str

class FlashcardDeckList(BaseModel):
    title: str
    cards: list[FlashcardItem]

class QuizQuestionItem(BaseModel):
    question_text: str
    option_a: str
    option_b: str
    option_c: str
    option_d: str
    correct_option: str
    explanation: str

class QuizList(BaseModel):
    title: str
    questions: list[QuizQuestionItem]

def generate_flashcards_ai(text_content: str, num_cards: int = 10):
    prompt = (
        f"You are a study helper. Based on this text, generate {num_cards} flashcards with front and back properties.\n\n"
        f"Text Content:\n{text_content[:15000]}"
    )
    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=FlashcardDeckList,
        ),
    )
    return json.loads(response.text)

def generate_quiz_ai(text_content: str, num_questions: int = 5):
    prompt = (
        f"You are a study helper. Based on this text, generate a multiple-choice quiz of {num_questions} questions.\n"
        "Each question must have 4 options (a, b, c, d), the correct option ('A', 'B', 'C', or 'D'), and a helpful explanation.\n\n"
        f"Text Content:\n{text_content[:15000]}"
    )
    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=QuizList,
        ),
    )
    return json.loads(response.text)